from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Rebuilding_Town_Final_Report_Word.docx"
OUTPUT = ROOT / "Rebuilding_Town_Final_Report_Word_Scene_Focused.docx"
DIAGRAM = ROOT / "report_assets" / "scene_structure_white.png"
FONT = Path("C:/Windows/Fonts/malgun.ttf")
FONT_BOLD = Path("C:/Windows/Fonts/malgunbd.ttf")


def font(size, bold=False):
    return ImageFont.truetype(str(FONT_BOLD if bold else FONT), size)


def draw_box(draw, xy, title, lines):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill="#FFFFFF", outline="#3D4A46", width=4)
    draw.rectangle((x1, y1, x2, y1 + 58), fill="#DDE8DD", outline="#3D4A46", width=3)
    draw.text(((x1 + x2) / 2, y1 + 29), title, font=font(25, True),
              fill="#111111", anchor="mm")
    y = y1 + 88
    for line in lines:
        draw.text(((x1 + x2) / 2, y), line, font=font(18),
                  fill="#242424", anchor="mm")
        y += 29


def arrow(draw, start, end):
    draw.line((start, end), fill="#596A63", width=4)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) > abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 16 * direction, ey - 9), (ex - 16 * direction, ey + 9)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 9, ey - 16 * direction), (ex + 9, ey - 16 * direction)]
    draw.polygon(points, fill="#596A63")


def build_diagram():
    image = Image.new("RGB", (3000, 2050), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((1500, 55), "주요 씬 구성도", font=font(46, True),
              fill="#111111", anchor="mm")
    draw.text((1500, 105), "상자 제목은 씬 파일명, 본문은 대표 구성 요소",
              font=font(21), fill="#555555", anchor="mm")

    draw_box(draw, (1050, 155, 1950, 395), "main.tscn", [
        "DayNightCycle · GroundTileMap · Trees",
        "Buildings · Citizens · Animals · MobSpawner",
        "Player · Pickups · PauseUI · PlayerHUD",
    ])

    draw.text((85, 480), "캐릭터 및 몹 씬", font=font(22, True), fill="#3D4A46")
    row2 = [
        ((45, 520, 425, 745), "player.tscn",
         ["AnimatedSprite2D · Camera2D", "SwordHitbox · ThrustHitbox"]),
        ((465, 520, 845, 745), "citizen.tscn",
         ["BodySprite · CollisionShape2D", "citizen.gd"]),
        ((885, 520, 1265, 745), "deer.tscn",
         ["BodySprite · AttackEffect", "CollisionShape2D · HealthBar"]),
        ((1305, 520, 1685, 745), "chicken.tscn",
         ["BodySprite · AttackEffect", "CollisionShape2D · HealthBar"]),
        ((1725, 520, 2105, 745), "boar.tscn",
         ["BodySprite · AttackEffect", "CollisionShape2D · HealthBar"]),
        ((2145, 520, 2525, 745), "bear.tscn",
         ["BodySprite · AttackEffect", "CollisionShape2D · HealthBar"]),
        ((2565, 520, 2945, 745), "tiger.tscn",
         ["BodySprite · AttackEffect", "CollisionShape2D · HealthBar"]),
    ]
    for box, title, lines in row2:
        draw_box(draw, box, title, lines)

    draw.text((85, 860), "건물 씬", font=font(22, True), fill="#3D4A46")
    row3 = [
        ((80, 900, 610, 1135), "mine_building.tscn",
         ["Stages 0~3 · InteractionArea", "PromptAnchor"]),
        ((670, 900, 1200, 1135), "blacksmith_building.tscn",
         ["Stages 0~3 · InteractionArea", "PromptAnchor"]),
        ((1260, 900, 1790, 1135), "gwana_building.tscn",
         ["Stages 0~4 · InteractionArea", "PromptAnchor"]),
        ((1850, 900, 2440, 1135), "restoration_lab_building.tscn",
         ["Stages 0~3 · InteractionArea", "RestorationProgress"]),
        ((2500, 900, 2920, 1135), "defense_wall.tscn",
         ["StageSprite · wall state", "defense_wall.gd"]),
    ]
    for box, title, lines in row3:
        draw_box(draw, box, title, lines)

    draw.text((85, 1260), "아이템 및 UI 씬", font=font(22, True), fill="#3D4A46")
    row4 = [
        ((80, 1300, 610, 1540), "artifact_item.tscn",
         ["Shadow · Sprite2D", "CollisionShape2D · PickupArea"]),
        ((670, 1300, 1200, 1540), "player_hud.tscn",
         ["DayLabel · HealthBar · CoinPanel", "TownStatus · ArtifactSlots"]),
        ((1260, 1300, 1790, 1540), "artifact_catalog_ui.tscn",
         ["ItemList · ArtifactImage", "Effect · Description"]),
        ((1850, 1300, 2380, 1540), "pause_ui.tscn",
         ["Overlay · Menu · Controls", "Resume · Save · Quit"]),
        ((2440, 1300, 2920, 1540), "game_over_ui.tscn",
         ["ScoreList · TotalScore", "Restart · Quit"]),
    ]
    for box, title, lines in row4:
        draw_box(draw, box, title, lines)

    draw.text((265, 1670), "환경 및 보조 씬", font=font(22, True), fill="#3D4A46")
    row5 = [
        ((260, 1710, 860, 1940), "coin.tscn",
         ["Sprite2D · CollisionShape2D", "coin.gd"]),
        ((920, 1710, 1520, 1940), "torch.tscn",
         ["불꽃 Sprite · Light2D", "torch.gd"]),
        ((1580, 1710, 2180, 1940), "e_key.tscn",
         ["E키 Sprite · 표시 제어", "e_key.gd"]),
        ((2240, 1710, 2840, 1940), "tree_zelkova_variant01.tscn",
         ["Sprite2D", "환경 배치용 변형 씬"]),
    ]
    for box, title, lines in row5:
        draw_box(draw, box, title, lines)

    image.save(DIAGRAM)


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_paragraph(doc, exact):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact:
            return paragraph
    raise ValueError(exact)


def remove_range(doc, start_text, end_text):
    start = find_paragraph(doc, start_text)._p
    end = find_paragraph(doc, end_text)._p
    current = start
    while current is not None and current is not end:
        nxt = current.getnext()
        remove_element(current)
        current = nxt


def add_before(doc, anchor, text, style=None, size=None, bold=False):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "맑은 고딕"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "맑은 고딕")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    anchor._p.addprevious(paragraph._p)
    return paragraph


def replace_text_everywhere(doc, old, new):
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        paragraph.text = paragraph.text.replace(old, new)


def remove_table_row_containing(doc, needle):
    for table in doc.tables:
        for row in list(table.rows):
            if needle in " ".join(cell.text for cell in row.cells):
                table._tbl.remove(row._tr)


def main():
    build_diagram()
    doc = Document(SOURCE)

    # Replace the original dark architecture image.
    image_paragraphs = [p for p in doc.paragraphs if p._p.xpath(".//a:blip")]
    architecture_paragraph = image_paragraphs[2]
    architecture_paragraph.clear()
    architecture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    architecture_paragraph.add_run().add_picture(str(DIAGRAM), width=Inches(6.5))
    find_paragraph(doc, "그림 2. 핵심 씬과 전역 시스템의 연결 구조").text = (
        "그림 2. 주요 씬 파일과 대표 구성 요소"
    )

    # Remove file-extension composition and the old scene summary table.
    remove_range(doc, "- 2.2 사용 파일 구성", "- 2.4 UI/UX 설계")

    # Insert scene-focused explanations immediately below the diagram intro.
    anchor = find_paragraph(doc, "- 2.4 UI/UX 설계")
    add_before(doc, anchor, "- 2.2 씬 상세 설명", "Heading 2")

    details = [
        ("main.tscn", "Heading 3",
         "게임 전체를 조립하는 최상위 씬이다. MountainBackground와 GroundTileMap이 배경과 지형을 구성하고, "
         "DayNightCycle이 시간 흐름을 담당한다. Buildings, DefenseWalls, Citizens, Animals, Pickups는 기능별 "
         "컨테이너이며 Player와 네 종류의 UI 씬이 직접 인스턴스로 배치된다. main.gd는 날짜 증가, 광산 수입, "
         "주민 보충, 습격 호출, 건물 효과, 승리와 결과 화면을 연결한다."),
        ("player.tscn", "Heading 3",
         "Player CharacterBody2D 아래에 이동 충돌을 위한 CollisionShape2D와 캐릭터 표현용 AnimatedSprite2D가 있다. "
         "SwordTrail과 SwordHitbox는 기본 연속 공격을, ThrustEffect와 ThrustHitbox는 찌르기 공격을 표현하고 판정한다. "
         "Camera2D는 플레이어 이동을 따라가며, player.gd는 체력·동전·공격력·유물 슬롯 2개와 입력 처리를 관리한다."),
        ("citizen.tscn", "Heading 3",
         "Citizen CharacterBody2D, BodySprite, CollisionShape2D로 구성된 단순한 주민 씬이다. citizen.gd가 배회, 적대 몹 "
         "감지, 도망, 피격과 사망을 처리한다. main.tscn의 CitizenSpawnPoints 여덟 곳에 생성되며 관아 단계에 따라 "
         "수용 인원이 증가하고 결원이 있으면 하루에 한 명씩 보충된다."),
        ("deer.tscn · chicken.tscn · boar.tscn · bear.tscn · tiger.tscn", "Heading 3",
         "다섯 동물 씬은 공통으로 CharacterBody2D, BodySprite, AttackEffect, CollisionShape2D, HealthBar를 사용하고 "
         "animal.gd를 공유한다. 사슴과 닭은 비적대 개체이며 멧돼지·곰·호랑이는 플레이어와 주민을 공격한다. 각 씬은 "
         "체력, 이동 속도, 공격력, 동전 드롭 수만 다르게 설정해 공통 AI를 재사용한다. MobSpawner가 일반 개체와 "
         "3일 주기 습격 개체를 실행 중 생성한다."),
        ("mine_building.tscn · blacksmith_building.tscn", "Heading 3",
         "두 건물은 Stages 아래에 Stage0~Stage3 Sprite2D를 두고 복구 단계에 맞는 외형만 표시한다. InteractionArea는 "
         "플레이어 접근을 감지하고 PromptAnchor는 E키 표시 위치를 제공한다. upgrade_building.gd가 업그레이드 비용과 "
         "건설 연출을 공통 처리하며, 광산은 일일 동전을 생산하고 대장간은 플레이어 공격력을 높인다."),
        ("gwana_building.tscn · restoration_lab_building.tscn", "Heading 3",
         "관아는 Stage0~Stage4와 InteractionArea, PromptAnchor로 구성되며 최종 4단계가 승리 조건에 포함된다. 관아 단계는 "
         "주민 수용 인원을 단계당 두 명씩 늘린다. 유물 복원소는 Stage0~Stage3 외에 RestorationProgress를 포함하며, "
         "restoration_lab.gd가 플레이어 슬롯의 손상 유물을 받아 복원 시간과 비용을 계산하고 ArtifactCatalog에 등록한다."),
        ("defense_wall.tscn", "Heading 3",
         "마을 건물 사이에 배치되는 방어벽 씬이다. defense_wall.gd가 양옆 건물의 복구 단계를 확인해 벽의 단계를 결정한다. "
         "현재는 건물 성장 상태를 외형으로 보여 주는 역할이 중심이며, 향후 적 이동 지연이나 피해 감소 기능을 연결할 수 있다."),
        ("artifact_item.tscn과 artifacts/items/", "Heading 3",
         "artifact_item.tscn은 Shadow, Sprite2D, CollisionShape2D, PickupArea로 구성된 손상 유물의 공통 부모다. "
         "artifacts/items 폴더의 유물 씬 10종은 이 씬을 상속하고 artifact_id와 텍스처만 지정한다. artifact_item.gd는 "
         "플레이어 접근과 슬롯 전달을 처리하며, animal.gd는 아직 복원되지 않은 유물 중 하나를 드롭 대상으로 선택한다."),
        ("player_hud.tscn", "Heading 3",
         "DayLabel, HealthBar, CoinPanel, TownStatusPanel, ArtifactSlots로 구성된다. 플레이 중 항상 필요한 날짜·체력·동전· "
         "건물·주민 상태를 보여 주며, 두 개의 유물 슬롯 아이콘과 중앙 안내 메시지를 갱신한다."),
        ("artifact_catalog_ui.tscn · pause_ui.tscn · game_over_ui.tscn", "Heading 3",
         "ArtifactCatalogUI는 왼쪽 ItemList와 오른쪽 이미지·효과·설명 영역으로 도감을 구성한다. PauseUI는 Overlay와 Menu "
         "안에 조작 설명, 계속, 저장, 종료 버튼을 배치한다. GameOverUI는 건물·주민·동전·유물·날짜 점수를 ScoreList에 "
         "표시하고 TotalScore, 재시작, 종료 버튼을 제공한다."),
        ("coin.tscn · torch.tscn · e_key.tscn · 환경 씬", "Heading 3",
         "coin.tscn은 필드 재화의 표시와 획득 충돌을 담당한다. torch.tscn은 불꽃 스프라이트와 조명을 묶은 환경 오브젝트이며, "
         "e_key.tscn은 건물 상호작용 가능 상태를 표시한다. tree_zelkova_variant01~03과 bush_grass_variant01~05는 "
         "Sprite2D 중심의 배치용 씬으로 월드의 반복 장식을 분리한다."),
    ]
    for title, style, body in details:
        add_before(doc, anchor, title, style)
        add_before(doc, anchor, body, "Normal")

    # Renumber the UI and system sections after the removed section.
    anchor.text = "- 2.3 UI/UX 설계"
    find_paragraph(doc, "- 2.5 세부 시스템 설계").text = "- 2.4 세부 시스템 설계"

    # Remove the complete test chapter, including its tables and callout.
    remove_range(doc, "3. 테스트 및 검증", "4. 개발 중 장애요인과 문제 해결 과정")
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == "테스트 및 검증":
            remove_element(paragraph._p)

    # Renumber following chapters.
    replace_text_everywhere(doc, "4. 개발 중 장애요인과 문제 해결 과정",
                            "3. 개발 중 장애요인과 문제 해결 과정")
    replace_text_everywhere(doc, "- 4.1 구조상 남은 위험", "- 3.1 구조상 남은 위험")
    replace_text_everywhere(doc, "5. 성과 및 결론", "4. 성과 및 결론")
    replace_text_everywhere(doc, "- 5.1 기대효과 및 활용방안", "- 4.1 기대효과 및 활용방안")
    replace_text_everywhere(doc, "- 5.2 프로젝트 성과", "- 4.2 프로젝트 성과")
    replace_text_everywhere(doc, "- 5.3 후기", "- 4.3 후기")

    # Remove remaining test-oriented statements outside the deleted chapter.
    replace_text_everywhere(
        doc,
        "Godot 씬 상속, Autoload, 충돌 레이어, 자동 테스트를 설명하는 교육 자료로 활용 가능하다.",
        "Godot 씬 구성, Autoload, 충돌 레이어와 공통 스크립트 구조를 설명하는 교육 자료로 활용 가능하다.",
    )
    replace_text_everywhere(
        doc,
        "전투·충돌·습격·운영·승리 자동 테스트 코드 마련",
        "전투·충돌·습격·운영·승리 기능을 파일별로 분리",
    )
    replace_text_everywhere(
        doc,
        "41개 tscn 씬, 공통 스크립트와 자동 테스트 구성",
        "41개 tscn 씬, 공통 스크립트와 이미지 리소스 구성",
    )
    replace_text_everywhere(doc, "통과 테스트", "충돌 설정 확인")
    replace_text_everywhere(doc, "습격 테스트", "생성 조건 확인")
    replace_text_everywhere(doc, "운영 테스트", "운영 로직 확인")
    replace_text_everywhere(doc, "승리 테스트", "승리 조건 확인")
    replace_text_everywhere(
        doc,
        "향후에는 실제 플레이 테스트 데이터를 수집해",
        "향후에는 실제 플레이 데이터를 수집해",
    )
    replace_text_everywhere(
        doc,
        "실제 플레이 스크린샷, Godot 4.6 자동 테스트 통과 로그를 추가해야 한다.",
        "실제 플레이 스크린샷과 실행 영상 자료를 추가해야 한다.",
    )
    remove_table_row_containing(doc, "자동 테스트 5종")

    # Update TOC wording and appendix summary.
    replace_text_everywhere(doc, "부록: 파일·씬 목록과 제출 점검표",
                            "부록: 파일 구조와 제출 점검표")
    replace_text_everywhere(
        doc,
        "전체 41개 씬의 상세 연결 관계는 프로젝트 루트의 `게임_씬_구조도.svg`와 "
        "`게임_전체_구조_분석.md`에서 확인할 수 있다.",
        "주요 씬의 구성 요소와 역할은 본문 2.2절에 정리했으며, 프로젝트 루트의 "
        "`게임_씬_구조도.svg`에는 전체 씬 연결 관계가 별도로 정리되어 있다.",
    )
    replace_text_everywhere(
        doc,
        "└─ tools/  자동 테스트·제작 보조 도구",
        "└─ tools/  이미지·리소스 제작 보조 도구",
    )

    doc.core_properties.title = "마을 재건 프로젝트 보고서 - 씬 중심 설명"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
