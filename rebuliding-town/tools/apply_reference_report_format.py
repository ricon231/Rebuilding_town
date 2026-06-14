from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "마을_재건_프로젝트_보고서.docx"
OUTPUT = ROOT / "마을_재건_프로젝트_보고서_객체지향양식.docx"

BLACK = RGBColor(0, 0, 0)
LIGHT_GRAY = "E7E6E6"
WHITE = "FFFFFF"


def set_run_font(run, name, size, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)


def set_style_font(style, name, size, bold=False):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = BLACK
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def remove_header_footer(section):
    for container in (section.header, section.footer):
        for paragraph in container.paragraphs:
            paragraph.clear()


def format_cover(doc):
    # Reference document: centered 28 pt headline, right-aligned team line,
    # left-aligned author/link information, otherwise restrained whitespace.
    cover = doc.paragraphs[:7]
    if len(cover) < 5:
        return

    cover[0].text = "객체지향 프로그래밍 최종보고서"
    cover[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover[0].paragraph_format.space_before = Pt(72)
    cover[0].paragraph_format.space_after = Pt(26)
    set_run_font(cover[0].runs[0], "HY헤드라인M", 22)

    cover[1].text = "<마을 재건>"
    cover[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover[1].paragraph_format.space_before = Pt(6)
    cover[1].paragraph_format.space_after = Pt(14)
    set_run_font(cover[1].runs[0], "HY헤드라인M", 28, True)

    cover[2].text = "팀 : [팀명 입력]"
    cover[2].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cover[2].paragraph_format.space_after = Pt(28)
    set_run_font(cover[2].runs[0], "HY견고딕", 16)

    # Keep the game-loop image but give it the same image-centered rhythm.
    cover[3].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover[3].paragraph_format.space_before = Pt(8)
    cover[3].paragraph_format.space_after = Pt(26)

    cover[4].text = "팀장 : [이름(학번) 입력]\n팀원 : [이름(학번) 입력]"
    cover[4].alignment = WD_ALIGN_PARAGRAPH.LEFT
    cover[4].paragraph_format.space_before = Pt(18)
    cover[4].paragraph_format.space_after = Pt(18)
    for run in cover[4].runs:
        set_run_font(run, "HY견고딕", 14)

    # The first table was a colored callout. Convert it to the reference's
    # plain source/video link block.
    if doc.tables:
        first = doc.tables[0]
        first.cell(0, 0).text = (
            "소스 코드 : [GitHub 주소 입력]\n"
            "시연 동영상 : [YouTube 주소 입력]"
        )
        shade(first.cell(0, 0), WHITE)
        for paragraph in first.cell(0, 0).paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_run_font(run, "맑은 고딕", 11)


def format_paragraphs(doc):
    for paragraph in doc.paragraphs:
        style = paragraph.style.name
        paragraph.paragraph_format.line_spacing = 1.15

        if style == "Heading 1":
            paragraph.paragraph_format.space_before = Pt(18)
            paragraph.paragraph_format.space_after = Pt(10)
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                set_run_font(run, "HY헤드라인M", 18)
        elif style == "Heading 2":
            paragraph.paragraph_format.space_before = Pt(13)
            paragraph.paragraph_format.space_after = Pt(7)
            paragraph.paragraph_format.keep_with_next = True
            if paragraph.text and not paragraph.text.lstrip().startswith("-"):
                text = paragraph.text
                paragraph.clear()
                run = paragraph.add_run("- " + text)
                set_run_font(run, "맑은 고딕", 16)
            else:
                for run in paragraph.runs:
                    set_run_font(run, "맑은 고딕", 16)
        elif style == "Heading 3":
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                set_run_font(run, "맑은 고딕", 14)
        elif style in ("List Bullet", "List Number"):
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.left_indent = Inches(0.35)
            for run in paragraph.runs:
                set_run_font(run, "맑은 고딕", 12)
        else:
            paragraph.paragraph_format.space_after = Pt(5)
            for run in paragraph.runs:
                if run.text.strip():
                    set_run_font(run, "맑은 고딕", 12)

    # Caption convention in the reference is simple, centered body text.
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("그림 "):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, "맑은 고딕", 10)
                run.italic = False


def format_tables(doc):
    for table_index, table in enumerate(doc.tables):
        table.autofit = False
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                shade(cell, LIGHT_GRAY if row_index == 0 and table_index > 0 else WHITE)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        set_run_font(
                            run,
                            "맑은 고딕",
                            10 if table_index > 0 else 11,
                            bold=(row_index == 0 and table_index > 0),
                        )


def main():
    doc = Document(SOURCE)
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(1.18)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)
        remove_header_footer(section)

    normal = doc.styles["Normal"]
    set_style_font(normal, "맑은 고딕", 12)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15
    set_style_font(doc.styles["Heading 1"], "HY헤드라인M", 18)
    set_style_font(doc.styles["Heading 2"], "맑은 고딕", 16)
    set_style_font(doc.styles["Heading 3"], "맑은 고딕", 14)
    set_style_font(doc.styles["List Bullet"], "맑은 고딕", 12)
    set_style_font(doc.styles["List Number"], "맑은 고딕", 12)

    format_paragraphs(doc)
    format_tables(doc)
    format_cover(doc)

    doc.core_properties.title = "마을 재건 객체지향 프로그래밍 최종보고서"
    doc.core_properties.subject = "객체지향 최종보고서 디자인 포맷 적용본"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
