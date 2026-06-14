from pathlib import Path
from datetime import date
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "마을_재건_프로젝트_보고서.docx"
ASSET_DIR = ROOT / "report_assets"
ASSET_DIR.mkdir(exist_ok=True)

FONT = Path("C:/Windows/Fonts/malgun.ttf")
FONT_BOLD = Path("C:/Windows/Fonts/malgunbd.ttf")
NAVY = "183A37"
GREEN = "4E7C59"
LIGHT_GREEN = "DCE9DC"
PALE = "F3F7F3"
GRAY = "606B67"
LIGHT_GRAY = "E7ECE9"
RED = "A33A35"
WHITE = "FFFFFF"
BLACK = "18201D"


def pil_font(size, bold=False):
    path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT
    return ImageFont.truetype(str(path), size)


def wrap(draw, text, font, max_width):
    words = text.split()
    lines, current = [], ""
    for word in words:
        trial = word if not current else current + " " + word
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def rounded_box(draw, box, title, lines, header_fill=LIGHT_GREEN, body_fill="172024"):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill="#" + body_fill, outline="#8CA398", width=3)
    draw.rectangle((x1, y1, x2, y1 + 58), fill="#" + header_fill)
    draw.text(((x1 + x2) / 2, y1 + 29), title, font=pil_font(25, True),
              fill="#111714", anchor="mm")
    y = y1 + 88
    for line in lines:
        draw.text(((x1 + x2) / 2, y), line, font=pil_font(19),
                  fill="#EDF3EF", anchor="mm")
        y += 31


def arrow(draw, start, end, color="#8CA398", width=5):
    draw.line((start, end), fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - 18 * direction, ey - 10), (ex - 18 * direction, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 10, ey - 18 * direction), (ex + 10, ey - 18 * direction)]
    draw.polygon(pts, fill=color)


def build_architecture_png():
    img = Image.new("RGB", (2400, 1500), "#080B0D")
    d = ImageDraw.Draw(img)
    d.text((1200, 52), "마을 재건 프로그램 구성도", font=pil_font(42, True),
           fill="#F4F7F5", anchor="mm")
    rounded_box(d, (830, 110, 1570, 310), "main.tscn",
                ["날짜·마을 운영·습격·승리 총괄", "main.gd / DayNightCycle / MobSpawner"],
                header_fill="F0D7A4")
    rounded_box(d, (90, 430, 500, 650), "player.tscn",
                ["이동·전투·체력·동전", "유물 슬롯 2개"])
    rounded_box(d, (560, 430, 970, 650), "citizen.tscn",
                ["최대 8명·위협 회피", "매일 결원 1명 보충"])
    rounded_box(d, (1030, 430, 1440, 650), "동물 씬 5종",
                ["deer / chicken", "boar / bear / tiger"])
    rounded_box(d, (1500, 430, 1910, 650), "건물 씬 5종",
                ["광산·대장간·관아", "복원소·방어벽"])
    rounded_box(d, (1970, 430, 2310, 650), "UI 씬 4종",
                ["HUD·도감", "일시정지·결과"])
    rounded_box(d, (310, 820, 800, 1040), "artifact_item.tscn",
                ["유물 공통 부모", "손상 유물 획득·운반"])
    rounded_box(d, (950, 820, 1450, 1040), "유물 씬 10종",
                ["개별 ID·텍스처", "animal.gd에서 동적 드롭"])
    rounded_box(d, (1600, 820, 2090, 1040), "ArtifactCatalog",
                ["전역 등록·효과·저장", "10종 완성 여부 제공"],
                header_fill="F0D7A4")
    rounded_box(d, (450, 1190, 1000, 1390), "환경·보조 씬",
                ["나무 3종·횃불·E키", "덤불 5종·농장·F키(미연결)"])
    rounded_box(d, (1400, 1190, 1950, 1390), "project.godot",
                ["실행 씬·Autoload·입력", "1600×900 / D3D12 / Godot 4.6"],
                header_fill="F0D7A4")
    arrow(d, (1200, 310), (295, 430))
    arrow(d, (1200, 310), (765, 430))
    arrow(d, (1200, 310), (1235, 430))
    arrow(d, (1200, 310), (1705, 430))
    arrow(d, (1200, 310), (2140, 430))
    arrow(d, (1235, 650), (1200, 820), color="#E2BD73")
    arrow(d, (800, 930), (950, 930), color="#78B0D3")
    arrow(d, (1450, 930), (1600, 930))
    arrow(d, (1675, 1190), (1675, 1040))
    img.save(ASSET_DIR / "architecture.png")


def build_game_loop_png():
    img = Image.new("RGB", (2200, 820), "#F7F9F7")
    d = ImageDraw.Draw(img)
    d.text((1100, 65), "핵심 플레이 순환 구조", font=pil_font(42, True),
           fill="#" + NAVY, anchor="mm")
    nodes = [
        ("탐색·전투", "몹 처치와 생존"),
        ("자원 획득", "동전·손상 유물"),
        ("마을 재건", "건물 단계 상승"),
        ("유물 복원", "효과·도감 등록"),
        ("습격 방어", "3일 주기 방어"),
    ]
    xs = [210, 650, 1090, 1530, 1970]
    for x, (title, detail) in zip(xs, nodes):
        d.rounded_rectangle((x - 175, 235, x + 175, 485), radius=24,
                            fill="#" + LIGHT_GREEN, outline="#" + GREEN, width=4)
        d.text((x, 310), title, font=pil_font(29, True), fill="#" + NAVY, anchor="mm")
        d.text((x, 390), detail, font=pil_font(21), fill="#" + BLACK, anchor="mm")
    for left, right in zip(xs[:-1], xs[1:]):
        arrow(d, (left + 180, 360), (right - 185, 360), color="#" + GREEN, width=6)
    d.arc((240, 500, 1970, 760), start=5, end=175, fill="#" + GREEN, width=6)
    d.polygon([(245, 628), (275, 610), (272, 646)], fill="#" + GREEN)
    d.text((1100, 700), "성장한 전투력과 마을 운영 능력으로 다음 날짜의 위협에 대응",
           font=pil_font(23, True), fill="#" + GRAY, anchor="mm")
    img.save(ASSET_DIR / "game_loop.png")


def build_ui_png():
    img = Image.new("RGB", (2200, 1240), "#10191A")
    d = ImageDraw.Draw(img)
    d.text((1100, 52), "main.tscn 기반 실제 화면 구조 배치도", font=pil_font(40, True),
           fill="#F3F7F3", anchor="mm")
    d.rounded_rectangle((80, 110, 2120, 1150), radius=25, fill="#243536",
                        outline="#91A69B", width=4)
    d.rectangle((100, 130, 2100, 965), fill="#9CC4CF")
    d.rectangle((100, 680, 2100, 965), fill="#7C9C66")
    d.polygon([(100, 680), (420, 400), (700, 680)], fill="#5B7E72")
    d.polygon([(520, 680), (900, 330), (1260, 680)], fill="#527067")
    d.polygon([(1150, 680), (1570, 380), (2020, 680)], fill="#58776A")

    # HUD
    d.rounded_rectangle((135, 160, 620, 320), radius=15, fill="#172024CC",
                        outline="#DCE9DC", width=3)
    d.text((165, 190), "PlayerHUD", font=pil_font(25, True), fill="#F4F7F5")
    d.text((165, 235), "체력  |  동전  |  현재 날짜", font=pil_font(22), fill="#F4F7F5")
    d.text((165, 275), "마을 건물·주민 상태", font=pil_font(22), fill="#F4F7F5")
    d.rounded_rectangle((1540, 160, 2055, 300), radius=15, fill="#172024CC",
                        outline="#DCE9DC", width=3)
    d.text((1797, 195), "유물 슬롯 2칸", font=pil_font(25, True),
           fill="#F4F7F5", anchor="mm")
    d.rectangle((1610, 230, 1710, 275), outline="#DCE9DC", width=3)
    d.rectangle((1840, 230, 1940, 275), outline="#DCE9DC", width=3)

    # World objects
    for x, name in [(500, "광산"), (860, "대장간"), (1220, "관아"), (1580, "복원소")]:
        d.rounded_rectangle((x - 120, 680, x + 120, 880), radius=18,
                            fill="#8A654B", outline="#E7D6B6", width=4)
        d.text((x, 770), name, font=pil_font(25, True), fill="#FFF4DE", anchor="mm")
        d.text((x, 825), "E 상호작용", font=pil_font(18), fill="#FFF4DE", anchor="mm")
    d.ellipse((970, 565, 1050, 645), fill="#E9D8B0", outline="#202020", width=3)
    d.text((1010, 615), "P", font=pil_font(24, True), fill="#202020", anchor="mm")
    d.text((1010, 545), "플레이어", font=pil_font(20, True), fill="#FFFFFF", anchor="mm")
    for x in [350, 680, 1370, 1840]:
        d.ellipse((x - 28, 600, x + 28, 656), fill="#D0E4CC")
    for x in [250, 1930]:
        d.ellipse((x - 35, 545, x + 35, 615), fill="#9D4B42")
        d.text((x, 525), "적", font=pil_font(19, True), fill="#FFFFFF", anchor="mm")

    d.rounded_rectangle((650, 900, 1550, 955), radius=12, fill="#172024",
                        outline="#DCE9DC", width=2)
    d.text((1100, 928), "중앙 안내 메시지: 습격·건물 복구·유물 획득 상태",
           font=pil_font(21), fill="#F4F7F5", anchor="mm")
    d.rectangle((100, 985, 2100, 1125), fill="#172024")
    d.text((1100, 1020), "오버레이 UI", font=pil_font(24, True),
           fill="#DCE9DC", anchor="mm")
    d.text((1100, 1070), "ESC: PauseUI  |  C: ArtifactCatalogUI  |  사망/승리: GameOverUI",
           font=pil_font(22), fill="#F4F7F5", anchor="mm")
    d.text((1100, 1195), "※ 실행 캡처가 아니라 현재 씬과 UI 코드의 배치 관계를 시각화한 화면 구조도",
           font=pil_font(19), fill="#AEBDB5", anchor="mm")
    img.save(ASSET_DIR / "ui_layout.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + key))
        if node is None:
            node = OxmlElement("w:" + key)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Malgun Gothic"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "맑은 고딕")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Heading 1", 16, NAVY, 16, 8),
        ("Heading 2", 13, GREEN, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def add_heading(doc, text, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_para(doc, text="", bold_lead=None, align=None, color=None, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, color=color, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(item))


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_run_font(p.add_run(item))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(p.add_run(str(text)), size=9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_note(doc, title, text, color=GREEN):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE)
    p = cell.paragraphs[0]
    r = p.add_run(title + "  ")
    set_run_font(r, bold=True, color=color)
    set_run_font(p.add_run(text), color=BLACK)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_image(doc, path, width, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(8)
    set_run_font(cap.add_run(caption), size=9, italic=True, color=GRAY)


def page_break(doc):
    doc.add_page_break()


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    configure_styles(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("GAME PROJECT REPORT"), size=12, bold=True, color=GREEN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("마을 재건"), size=30, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Rebuilding Town 프로그램 분석 및 개발 보고서"),
                 size=16, color=GREEN)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSET_DIR / "game_loop.png"), width=Inches(6.1))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(f"작성 기준일  {date.today().isoformat()}"), size=11, color=GRAY)
    add_note(doc, "제출 필수 링크",
             "GitHub 소스코드 저장소: 현재 미등록 / YouTube 시연 영상: 현재 미등록. "
             "최종 제출 전 공개 GitHub 저장소와 YouTube 영상 주소를 반드시 입력해야 한다.",
             color=RED)

    page_break(doc)
    add_heading(doc, "목차", 1)
    add_numbered(doc, [
        "프로젝트 개요",
        "개발 프로그램 설명",
        "테스트 및 검증",
        "개발 중 장애요인과 문제 해결 과정",
        "성과 및 결론",
        "부록: 파일·씬 목록과 제출 점검표",
    ])
    add_heading(doc, "보고서 요약", 1)
    add_para(doc,
             "마을 재건은 전투로 자원을 확보하고 마을 건물을 복구하며, 손상된 전통 유물을 "
             "복원하는 2D 액션·운영 게임이다. 전투, 날짜 진행, 주민 관리, 건물 성장, 유물 "
             "수집을 하나의 반복 구조로 묶었으며, 유물 10종 복원과 관아 4단계 달성을 최종 "
             "목표로 설정하였다.")
    add_table(doc, ["구분", "내용"], [
        ["엔진", "Godot 4.6 / GDScript"],
        ["화면", "1600×900, 2D 횡스크롤 월드"],
        ["핵심 시스템", "전투, 날짜, 습격, 건물 재건, 주민, 유물 복원"],
        ["승리 조건", "유물 10종 전체 복원 + 관아 4단계"],
        ["프로젝트 규모", "41개 tscn 씬, 공통 스크립트와 자동 테스트 구성"],
    ], [2100, 7260])

    page_break(doc)
    add_heading(doc, "1. 프로젝트 개요", 1)
    add_heading(doc, "1.1 작품 주제 선정 배경", 2)
    add_para(doc,
             "일반적인 액션 게임은 전투 성과가 장비 수치에만 누적되는 경우가 많다. 본 프로젝트는 "
             "플레이어의 전투 결과가 마을의 재건, 주민의 증가, 문화유산의 복원으로 이어지도록 "
             "설계하여 행동의 의미를 시각적·운영적 변화로 보여 주는 것을 목표로 한다.")
    add_bullets(doc, [
        "폐허가 된 마을을 단계적으로 복구하는 명확한 성장 경험 제공",
        "전투와 경영을 분리하지 않고 동전·건물·주민·습격으로 연결",
        "전통 유물을 단순 수집품이 아니라 능력 효과와 승리 조건으로 활용",
        "짧은 조작 체계 안에서 액션과 운영 판단을 동시에 경험하도록 구성",
    ])
    add_heading(doc, "1.2 벤치마킹 자료", 2)
    add_table(doc, ["대상", "참고 요소", "프로젝트 반영 방식"], [
        ["Stardew Valley", "일상 진행과 거점 성장", "날짜 단위 수입과 마을 상태 변화"],
        ["Kingdom: Two Crowns", "거점 방어와 주기적 적의 압박", "3일 주기 오른쪽 습격과 주민 보호"],
        ["Moonlighter", "전투 보상과 거점 경제의 연결", "몹 처치 동전을 건물 복구 비용으로 사용"],
        ["박물관형 수집 콘텐츠", "유물 설명·도감·완성 목표", "손상 유물 획득, 복원, 카탈로그 등록"],
    ], [1900, 2600, 4860])
    add_note(doc, "벤치마킹 원칙",
             "특정 작품의 규칙을 그대로 복제하지 않고, 성장 피드백·거점 방어·수집 완성이라는 "
             "설계 원리를 프로젝트 규모에 맞게 단순화하였다.")
    add_image(doc, ASSET_DIR / "game_loop.png", 6.5, "그림 1. 전투와 마을 운영을 연결하는 핵심 게임 루프")

    add_heading(doc, "1.3 개발 환경 설명", 2)
    add_table(doc, ["분류", "구성"], [
        ["운영체제", "Windows 개발 환경"],
        ["하드웨어 기준", "x64 PC, Direct3D 12 지원 그래픽 장치, 1600×900 출력 환경"],
        ["게임 엔진", "Godot Engine 4.6"],
        ["언어", "GDScript, Godot tscn/tres 리소스"],
        ["개발 도구·IDE", "Godot Script Editor 및 프로젝트 파일 편집 도구"],
        ["렌더링", "Forward Plus, D3D12"],
        ["외부 서버", "게임 실행에 필요한 별도 서버 없음"],
        ["플랫폼", "현재 Windows 데스크톱 실행 기준"],
        ["전역 구성", "ArtifactCatalog Autoload"],
    ], [2200, 7160])
    add_para(doc,
             "프로젝트 설정상 3D 물리 엔진은 Jolt로 지정되어 있으나, 실제 게임 플레이는 "
             "CharacterBody2D, Area2D, TileMapLayer 중심의 2D 구조다.")

    page_break(doc)
    add_heading(doc, "2. 개발 프로그램 설명", 1)
    add_heading(doc, "2.1 전체 기능 설계", 2)
    add_image(doc, ASSET_DIR / "architecture.png", 6.6, "그림 2. 핵심 씬과 전역 시스템의 연결 구조")
    add_para(doc,
             "`main.tscn`은 월드의 최상위 씬이며 날짜, 건물 효과, 주민 보충, 습격, 승리와 "
             "게임오버를 조정한다. 플레이어·건물·UI·환경 씬은 메인에 직접 배치되고, 주민·몹· "
             "유물 아이템은 실행 중 조건에 따라 동적으로 생성된다.")
    add_heading(doc, "2.2 사용 파일 구성", 2)
    add_table(doc, ["확장자·유형", "역할", "대표 파일"], [
        [".tscn", "노드 구조, 씬 인스턴스, 충돌·스프라이트 설정", "main.tscn, player.tscn"],
        [".gd", "게임 규칙과 입력·AI·UI 로직", "main.gd, animal.gd"],
        [".tres", "애니메이션 프레임 등 재사용 리소스", "animal_*_frames.tres"],
        [".png / .svg", "캐릭터, 건물, 환경, 유물 이미지", "assets/, icon.svg"],
        ["project.godot", "메인 씬, Autoload, 렌더링·입력 설정", "project.godot"],
        ["tools/", "자동 테스트와 제작 보조 스크립트", "test_*.gd"],
    ], [1700, 3900, 3760])

    add_heading(doc, "2.3 주요 씬 구성", 2)
    add_table(doc, ["영역", "씬", "주요 책임"], [
        ["월드", "main.tscn", "전체 게임 진행과 하위 시스템 조율"],
        ["플레이어", "player.tscn", "이동, 공격, 체력, 동전, 유물 슬롯"],
        ["주민", "citizen.tscn", "마을 인구, 위협 회피, 사망 처리"],
        ["몹", "deer/chicken/boar/bear/tiger", "비적대·적대 AI와 드롭"],
        ["건물", "mine/blacksmith/gwana/restoration_lab", "경제·공격·인구·유물 복원"],
        ["방어", "defense_wall.tscn", "인접 건물 단계에 따른 외형 변화"],
        ["유물", "artifact_item + items 10종", "필드 획득, 운반, 복원 데이터"],
        ["UI", "HUD/catalog/pause/game_over", "상태 전달, 도감, 메뉴, 결과"],
        ["환경", "tree/torch/bush 등", "월드 장식과 조명"],
    ], [1350, 3600, 4410])

    page_break(doc)
    add_heading(doc, "2.4 UI/UX 설계", 2)
    add_image(doc, ASSET_DIR / "ui_layout.png", 6.6, "그림 3. main.tscn과 UI 코드에 근거한 화면 구조")
    add_heading(doc, "정보 배치 원칙", 3)
    add_bullets(doc, [
        "좌측 상단: 체력, 동전, 날짜와 마을 상태처럼 자주 확인하는 정보를 고정 배치",
        "우측 상단: 운반 중인 유물 슬롯 2개를 별도로 표시해 복원 행동을 유도",
        "화면 중앙: 플레이어, 주민, 몹, 건물을 같은 월드 공간에서 직접 확인",
        "중앙 하단: 습격·유물·건물 업그레이드 등 일시적 메시지 표시",
        "오버레이: ESC 일시정지, C 유물 도감, 승리·사망 시 결과 화면",
        "건물 근처: E키 상호작용 표시를 해당 대상 가까이에 배치",
    ])
    add_heading(doc, "사용자 경험 흐름", 3)
    add_numbered(doc, [
        "플레이어는 체력·동전·날짜를 확인하며 필드를 탐색한다.",
        "몹을 처치해 동전과 손상 유물을 얻는다.",
        "건물에 접근하면 E키 표시를 통해 업그레이드 또는 복원을 수행한다.",
        "C키 도감에서 복원된 유물과 효과를 확인한다.",
        "습격 메시지를 통해 마을 오른쪽에서 오는 적에 대비한다.",
        "목표 달성 또는 사망 시 결과 화면에서 성과를 확인한다.",
    ])
    add_note(doc, "화면 자료 제한",
             "현재 작업 환경에서 Godot 실행 파일을 확인할 수 없어 실제 플레이 스크린샷을 "
             "새로 촬영하지 못했다. 그림 3은 `main.tscn`, `player_hud.tscn`, UI 스크립트의 "
             "노드 배치와 표시 정보를 근거로 제작한 화면 구조도다.")

    add_heading(doc, "2.5 세부 시스템 설계", 2)
    add_heading(doc, "전투와 적 AI", 3)
    add_para(doc,
             "플레이어는 기본 공격 26을 기준으로 2연격과 찌르기를 사용한다. 적대 몹은 "
             "플레이어 또는 주민을 추적하며, 습격 모드에서는 우선 마을 중심을 향해 이동한다. "
             "몹과 주민의 충돌 마스크는 지형 레이어만 감지하도록 설정되어 개체끼리 밀어내지 않는다.")
    add_table(doc, ["몹", "성향", "체력", "속도", "동전"], [
        ["닭", "비적대", "20", "25", "1"],
        ["사슴", "비적대", "45", "50", "2"],
        ["멧돼지", "적대", "90", "55", "4"],
        ["곰", "적대", "160", "30", "7"],
        ["호랑이", "적대", "200", "70", "10"],
    ], [2100, 1800, 1700, 1700, 2060])
    add_heading(doc, "마을 운영", 3)
    add_table(doc, ["건물", "단계", "효과"], [
        ["광산", "0~3", "매일 단계 × 2 동전"],
        ["대장간", "0~3", "단계마다 공격력 +15%"],
        ["관아", "0~4", "단계 × 2명 수용, 최대 8명"],
        ["유물 복원소", "0~3", "단계 상승 시 복원 비용·시간 감소"],
        ["방어벽", "1~3", "인접 건물 단계에 따라 외형 변화"],
    ], [2400, 1500, 5460])
    add_heading(doc, "날짜와 습격", 3)
    add_bullets(doc, [
        "하루 약 180초, 일반 몹 생성 간격 약 14초",
        "일반 생성 개체 상한 8마리",
        "비적대 몹 확률은 첫날 70%에서 매일 5%p 감소하며 최소 15%",
        "호랑이는 5일차부터 등장",
        "3일마다 오른쪽에서 습격, 30일차까지 진행",
        "3마리로 시작해 회차마다 2마리 증가하며 최대 15마리",
    ])

    page_break(doc)
    add_heading(doc, "3. 테스트 및 검증", 1)
    add_para(doc,
             "프로젝트의 주요 규칙은 `tools/` 폴더의 Godot SceneTree 테스트로 검증하도록 구성되어 "
             "있다. 테스트는 실제 씬을 인스턴스화하고 속성, 충돌 레이어, 생성 수, 승리 화면을 "
             "assert로 확인한다.")
    add_table(doc, ["테스트 파일", "검증 내용", "기대 결과"], [
        ["test_combat_balance.gd", "플레이어 공격·피격 보호, 적 체력·공격·드롭", "설정 수치 일치"],
        ["test_mob_passthrough.gd", "플레이어·주민·몹 충돌 레이어", "개체끼리 밀지 않고 통과"],
        ["test_raid_flow.gd", "3일차 습격, 오른쪽 생성, 30일 상한", "3마리 시작·30일 15마리"],
        ["test_town_management.gd", "광산 수입, 대장간 보너스, 주민 수", "최대 8명·매일 1명 보충"],
        ["test_victory_condition.gd", "유물 10종과 관아 4단계", "두 조건 충족 시 승리 UI"],
    ], [2800, 3700, 2860])
    add_heading(doc, "3.1 정적 검증 결과", 2)
    add_bullets(doc, [
        "프로젝트의 tscn 씬 41개를 목록화하고 구조도에 모두 반영",
        "main.tscn의 직접 참조 씬과 스크립트 동적 생성 씬을 구분",
        "몹·주민 레이어 4 / 마스크 1, 플레이어 레이어 2 / 마스크 1 확인",
        "습격 조건: 3일 주기, 30일차 종료, 최대 15마리 코드 확인",
        "주민 상한 8명과 일일 1명 보충 로직 확인",
        "승리 조건: ArtifactCatalog 전체 등록과 관아 4단계 동시 검사 확인",
    ])
    add_note(doc, "실행 검증 상태",
             "자동 테스트 코드는 존재하지만 현재 보고서 생성 환경에서는 Godot 실행 파일이 "
             "검색되지 않아 이번 문서 작성 과정에서 테스트를 재실행하지 못했다. 최종 제출 전 "
             "Godot 4.6에서 각 test_*.gd를 headless 모드로 실행해 통과 로그를 첨부해야 한다.",
             color=RED)
    add_heading(doc, "3.2 최종 수동 테스트 시나리오", 2)
    add_table(doc, ["번호", "절차", "통과 기준"], [
        ["1", "플레이어와 몹·주민을 겹쳐 이동", "개체가 서로 밀지 않고 통과"],
        ["2", "3일차까지 진행", "오른쪽에서 적대 몹 3마리 생성"],
        ["3", "30일차와 33일차 습격 확인", "30일 15마리, 33일 추가 생성 없음"],
        ["4", "관아 4단계에서 주민 2명 사망", "다음 날부터 하루 1명씩 복구"],
        ["5", "손상 유물 획득 후 복원소 이용", "복원 완료 후 도감 등록·효과 반영"],
        ["6", "유물 10종 + 관아 4단계 달성", "승리 결과 화면 표시"],
    ], [700, 4760, 3900])

    page_break(doc)
    add_heading(doc, "4. 개발 중 장애요인과 문제 해결 과정", 1)
    add_table(doc, ["장애요인", "원인 분석", "해결 방법", "검증"], [
        ["몹끼리 접촉 시 밀림", "CharacterBody2D 상호 충돌", "레이어 4, 마스크 1로 지형만 감지", "통과 테스트"],
        ["검은 사각형 적 등장", "임시 적 씬·스프라이트 잔존", "관련 씬·스크립트·이미지 제거", "참조 검색"],
        ["튜토리얼 표시 혼선", "상태·입력 안내 중복", "튜토리얼 제거, 대상 근처 E키 표시 유지", "씬 참조 확인"],
        ["습격 규모가 불명확", "종료일·상한 부재", "3일 주기, 30일, 최대 15마리 명시", "습격 테스트"],
        ["주민 사망 후 회복 불가", "재생성 규칙 부재", "관아 수용력 내 매일 1명 보충", "운영 테스트"],
        ["승리 목표가 모호함", "수집과 건물 목표 분리", "유물 10종 + 관아 4단계 동시 조건", "승리 테스트"],
        ["UI 한국어 가독성", "문자 인코딩·폰트 의존성", "맑은 고딕 계열 사용과 UI 문구 일원화 권장", "실기기 확인 필요"],
    ], [1700, 2350, 3350, 1960])
    add_heading(doc, "4.1 구조상 남은 위험", 2)
    add_bullets(doc, [
        "main.gd가 날짜, 경제, 주민, 승리, 점수까지 담당해 기능 추가 시 복잡도가 빠르게 증가할 수 있다.",
        "방어벽은 현재 시각적 성장 요소 중심이므로 실제 방어 효과를 부여할 필요가 있다.",
        "30일 이후 정기 습격이 중단되어 미승리 플레이의 긴장감이 약해질 수 있다.",
        "bulid_farm.tscn, interact_f.tscn, 덤불 5종은 현재 실행 흐름에 연결되지 않는다.",
        "한국어 문자열은 운영체제와 편집기 인코딩을 UTF-8로 통일해 재검증할 필요가 있다.",
    ])

    page_break(doc)
    add_heading(doc, "5. 성과 및 결론", 1)
    add_heading(doc, "5.1 기대효과 및 활용방안", 2)
    add_bullets(doc, [
        "전통 유물을 게임 목표와 능력 효과로 연결해 문화유산에 대한 흥미를 높일 수 있다.",
        "전투와 마을 운영의 상호작용을 학습용 게임 기획 사례로 활용할 수 있다.",
        "Godot 씬 상속, Autoload, 충돌 레이어, 자동 테스트를 설명하는 교육 자료로 활용 가능하다.",
        "유물 설명, 지역 설화, 건물 역할을 확장하면 전시·박물관형 콘텐츠로 발전시킬 수 있다.",
        "농장, 방어벽 기능, 주민 직업을 추가해 장기 운영 게임으로 확장할 수 있다.",
    ])
    add_heading(doc, "5.2 프로젝트 성과", 2)
    add_table(doc, ["성과 영역", "구현 결과"], [
        ["게임 루프", "전투→자원→재건→복원→습격 방어의 순환 구조 완성"],
        ["콘텐츠", "몹 5종, 핵심 건물 5종, 유물 10종, UI 4종 구성"],
        ["마을 운영", "광산 수입, 대장간 강화, 관아 주민 관리, 복원소 성장"],
        ["난이도 구조", "날짜별 적대 비율, 호랑이 해금, 3일 주기 습격"],
        ["완료 조건", "유물 전체 복원과 관아 최종 단계의 복합 승리 목표"],
        ["품질 관리", "전투·충돌·습격·운영·승리 자동 테스트 코드 마련"],
    ], [2200, 7160])
    add_heading(doc, "5.3 후기", 2)
    add_para(doc,
             "이 프로젝트를 통해 하나의 기능을 추가하는 것보다 여러 시스템의 관계를 일관되게 "
             "만드는 일이 더 중요하다는 점을 확인할 수 있었다. 적의 충돌 방식 하나가 전투 감각과 "
             "주민 동선에 동시에 영향을 주고, 주민의 생성 규칙은 관아의 가치와 습격의 긴장도를 "
             "함께 바꾼다. 또한 승리 조건을 명확히 정의하면서 유물 수집과 마을 재건이 별개의 "
             "콘텐츠가 아니라 하나의 목표를 이루는 구조로 정리되었다.")
    add_para(doc,
             "향후에는 실제 플레이 테스트 데이터를 수집해 공격력, 드롭률, 복원 비용, 습격 간격을 "
             "정량적으로 조정하고, 방어벽과 주민 직업처럼 현재 시각적·기초 단계에 있는 요소를 "
             "게임 전략으로 확장하는 것이 중요하다.")
    add_note(doc, "결론",
             "마을 재건은 핵심 시스템이 서로 연결된 플레이 가능한 프로젝트 구조를 갖추었다. "
             "최종 제출 품질을 위해서는 GitHub 공개 저장소, YouTube 시연 영상, 실제 플레이 "
             "스크린샷, Godot 4.6 자동 테스트 통과 로그를 추가해야 한다.")

    page_break(doc)
    add_heading(doc, "부록 A. 파일 구조", 1)
    tree = (
        "rebuliding-town/\n"
        "├─ project.godot\n"
        "├─ main.tscn / main.gd\n"
        "├─ player.tscn / player.gd\n"
        "├─ citizen.tscn / citizen.gd\n"
        "├─ animal.gd + 동물 씬 5종\n"
        "├─ mob_spawner.gd\n"
        "├─ 건물 씬 5종 + 공통 건물 스크립트\n"
        "├─ artifact_item.tscn / artifact_catalog.gd\n"
        "├─ artifacts/items/  유물 씬 10종\n"
        "├─ UI 씬 4종\n"
        "├─ assets/  이미지·애니메이션 리소스\n"
        "└─ tools/  자동 테스트·제작 보조 도구"
    )
    p = doc.add_paragraph()
    set_cell = p.add_run(tree)
    set_run_font(set_cell, size=9.5, color=BLACK)
    set_cell.font.name = "D2Coding"
    set_cell._element.rPr.rFonts.set(qn("w:eastAsia"), "D2Coding")
    add_para(doc,
             "전체 41개 씬의 상세 연결 관계는 프로젝트 루트의 `게임_씬_구조도.svg`와 "
             "`게임_전체_구조_분석.md`에서 확인할 수 있다.")
    add_heading(doc, "부록 B. 제출 전 점검표", 1)
    add_table(doc, ["상태", "제출 항목"], [
        ["필수", "프로젝트를 공개 GitHub 저장소에 업로드하고 표지 URL 교체"],
        ["필수", "게임 시연 영상을 YouTube에 업로드하고 표지 URL 교체"],
        ["필수", "실제 게임 실행 화면 2~4장을 UI/UX 절에 삽입"],
        ["필수", "Godot 4.6에서 자동 테스트 5종 실행 후 통과 로그 첨부"],
        ["권장", "README에 실행 방법, 조작법, 승리 조건 명시"],
        ["권장", "미사용 씬의 유지·삭제 여부 정리"],
    ], [1200, 8160])

    # Header / footer
    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(hp.add_run("마을 재건 | 프로젝트 보고서"), size=8.5, color=GRAY)
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Rebuilding Town · 2026")
        set_run_font(run, size=8.5, color=GRAY)

    props = doc.core_properties
    props.title = "마을 재건 프로젝트 보고서"
    props.subject = "Godot 게임 프로그램 분석 및 개발 보고서"
    props.author = "Rebuilding Town Project"
    props.keywords = "Godot, 게임 개발, 마을 운영, 유물 복원"
    doc.save(OUT)


if __name__ == "__main__":
    build_architecture_png()
    build_game_loop_png()
    build_ui_png()
    build_docx()
    print(OUT)
